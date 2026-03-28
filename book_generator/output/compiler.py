"""Compile all approved chapters into .docx and .txt output files."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from book_generator.db.client import get_db
from book_generator.core.logger import logger

OUTPUT_DIR = Path("output_books")
OUTPUT_DIR.mkdir(exist_ok=True)


def _get_ordered_chapters(book_id: str) -> list[dict]:
    db = get_db()
    return (
        db.table("chapters")
        .select("chapter_number, title, content")
        .eq("book_id", book_id)
        .order("chapter_number")
        .execute()
        .data
    )


def compile_book(book_id: str, title: str) -> dict:
    """Generate .docx and .txt files. Returns dict with file paths."""
    chapters = _get_ordered_chapters(book_id)
    if not chapters:
        raise ValueError(f"No chapters found for book_id={book_id}")

    safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)[:60]
    docx_path = OUTPUT_DIR / f"{safe_title}.docx"
    txt_path = OUTPUT_DIR / f"{safe_title}.txt"

    # ── DOCX ──────────────────────────────────────────────────
    doc = Document()

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    doc.add_page_break()

    # Chapters
    for ch in chapters:
        heading = doc.add_heading(f"Chapter {ch['chapter_number']}: {ch['title']}", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        doc.add_paragraph(ch["content"] or "")
        doc.add_page_break()

    doc.save(str(docx_path))

    # ── TXT ───────────────────────────────────────────────────
    lines = [title.upper(), "=" * len(title), ""]
    for ch in chapters:
        lines.append(f"\nCHAPTER {ch['chapter_number']}: {ch['title'].upper()}")
        lines.append("-" * 60)
        lines.append(ch["content"] or "")
        lines.append("")
    txt_path.write_text("\n".join(lines), encoding="utf-8")

    # Persist paths to DB
    db = get_db()
    db.table("books").update({
        "output_docx": str(docx_path),
        "output_txt": str(txt_path),
        "status": "completed",
    }).eq("id", book_id).execute()

    logger.info(f"Book compiled: docx={docx_path} txt={txt_path}")
    return {"docx": str(docx_path), "txt": str(txt_path)}
